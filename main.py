from PIL import Image
from docx import Document
from docx.shared import Inches
import glob
import os

def sortKeyFunc(s):
    return int(os.path.basename(s)[:-4][7:])

imageList = []

for filename in sorted(glob.glob('C:/Users/user/Documents/FCA/Blueprints/img/*')):
    imageList.append(filename)

imageList.sort(key=sortKeyFunc)

#setup doc
document = Document()
tables = document.tables

while len(imageList) != 0:
    run = document.add_paragraph().add_run()
    print('adding 2 pictures')
    if len(imageList) != 0:
        run.add_picture(imageList.pop(0), width=Inches(3))
    if len(imageList) != 0:
        run.add_picture(imageList.pop(0), width=Inches(3))

#save the doc
document.save('C:/Users/user/Documents/FCA/Blueprints/L2_bp.docx')
print('done')
        

